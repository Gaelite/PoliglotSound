from ClassAudio import audio
from docx import Document

traductor = audio()
traducción = Document()
traductor.hablar("Habla para identificar tu idioma Speak to identify your language")
traductor.set_texto()
while True:
    traducción.add_heading((traductor.traductor.translate("Traducción", dest=traductor.texto_elegido).text),0)
    traductor.hablar(traductor.traductor.translate("¿A qué idioma quieres traducir? \nDi 'salir' para abandonar el traductor", dest=traductor.texto_elegido).text)
    print(traductor.traductor.translate("Chino mandarín \nEspañol\nInglés\nHindi\nÁrabe\nPortugués\nBengalí\nRuso\nJaponés\nPanyabí\nFrancés\nAlemán\nPuedes decir 'Salir' para abandonar el traductor", dest=traductor.texto_elegido).text)
    traductor.set_texto()
    if traductor.salir() == 1:
        traducción.add_paragraph((traductor.traductor.translate("Abandonó el traductor", dest=traductor.texto_elegido).text))
        break
    x = traductor.set_idioma_elegido()
    traducción.add_heading((traductor.traductor.translate(f"Idioma elegido: {x}", dest=traductor.texto_elegido).text),1)
    while True:
        traductor.hablar(traductor.traductor.translate("Dictar el texto a traducir\nDi 'salir' para abandonar el traductor", dest=traductor.texto_elegido).text)
        print(traductor.traductor.translate(f"traduciendo a {x}", dest=traductor.texto_elegido).text)
        traductor.set_texto()
        if traductor.salir() == 1:
            break
        traductor.hablar(traductor.traducir())
        print(traductor.get_texto(),"\n",traductor.traducir())
        traducción.add_paragraph(f"- {traductor.get_texto()}\n+ {traductor.traducir()}")

traducción.save(traductor.traductor.translate("Traducción", dest=traductor.texto_elegido).text)

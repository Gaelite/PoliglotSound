from flask import Flask,render_template
from ClassAudio import audio
from docx import Document

transc = Document()
transc.add_heading("Traducción",0)
traductor = audio()
app=Flask(__name__) 
app.static_folder = 'static'
idioma = "no seleccionado"
texto =  []

#decorador para asociar la ruta con la función 

@app.route("/") 

def index():
    return render_template("Trad.html",idioma = idioma)

@app.route('/CambiarTraduccion', methods=['POST'])
def CambiarTraduccion():
    global idioma
    idioma = traductor.set_idioma_elegido()
    texto.append(f"Traduciendo a {idioma}:")
    transc.add_heading(f"Idioma elegido: {idioma}",1)
    return render_template('Trad.html',idioma = idioma, texto_capturado=texto)

@app.route('/boton1', methods=['POST'])
def boton1():
    audio = traductor.set_texto()
    texto.append(f"-{audio}")
    transc.add_paragraph(f"- {audio}")
    audio = traductor.traducir()
    texto.append(f"+{audio}")
    transc.add_paragraph(f"+ {audio}")
    return render_template('Trad.html', texto_capturado=texto, idioma = idioma)

@app.route('/transcripcion', methods=['POST'])
def transcripcion():
    transc.save("Traducción")
    texto = []
    idioma = "No seleccionado"
    return render_template('Trad.html', texto_capturado=texto, idioma = idioma)

if __name__=='__main__': 

    app.run(debug=True) 